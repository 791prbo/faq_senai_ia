import time
import tiktoken
import pypdf
import docx
from g4f.client import Client
from groq import Groq
from config import MODELO_DEFAULT_G4F, MODELO_DEFAULT_GROQ
import io
import docx

def calcular_tokens(texto, modelo="gpt-4o-mini"):
    """Calcula a quantidade de tokens de uma string."""
    try:
        codificador = tiktoken.encoding_for_model(modelo)
    except KeyError:
        codificador = tiktoken.get_encoding("cl100k_base")
    return len(codificador.encode(texto))

def extrair_texto_de_arquivo(uploaded_file):
    """Extrai o texto de arquivos .txt, .pdf ou .docx."""
    if uploaded_file is None:
        return ""
    
    # 1. Arquivos de texto simples
    if uploaded_file.name.endswith('.txt'):
        return uploaded_file.getvalue().decode("utf-8")
    
    # 2. Arquivos do Word (.docx)
    elif uploaded_file.name.endswith('.docx'):
        doc = docx.Document(uploaded_file)
        texto_extraido = "\n".join([paragrafo.text for paragrafo in doc.paragraphs if paragrafo.text])
        return texto_extraido

    # 3. Arquivos PDF (.pdf)
    elif uploaded_file.name.endswith('.pdf'):
        leitor_pdf = pypdf.PdfReader(uploaded_file)
        texto_extraido = ""
        for pagina in leitor_pdf.pages:
            texto = pagina.extract_text()
            if texto:
                texto_extraido += texto + "\n"
        return texto_extraido
        
    return ""

def gerar_resposta_ia(provedor, mensagens, groq_key=None, temperatura=0.2):
    """Gera a resposta da IA e retorna o texto acompanhado do tempo de execução."""
    inicio = time.time()
    
    if provedor == "GPT-4o Mini (Via G4F)":
        client = Client()
        resposta = client.chat.completions.create(
            model=MODELO_DEFAULT_G4F,
            messages=mensagens,
            temperature=temperatura
        )
        texto = resposta.choices[0].message.content

    elif provedor == "Llama 3.3 70B (Via Groq)":
        if not groq_key:
            raise ValueError("Chave 'GROQ_API_KEY' não foi encontrada no secrets.toml.")
        
        client_groq = Groq(api_key=groq_key)
        resposta = client_groq.chat.completions.create(
            model=MODELO_DEFAULT_GROQ,
            messages=mensagens,
            temperature=temperatura
        )
        texto = resposta.choices[0].message.content
    else:
        raise ValueError("Provedor de IA desconhecido.")

    tempo_decorrido = round(time.time() - inicio, 2)
    return texto, tempo_decorrido

def gerar_template_word_bytes():
    """Gera o arquivo Word (.docx) padronizado em memória para download."""
    doc = docx.Document()
    
    # Título Principal
    doc.add_heading("BASE DE DADOS OFICIAL - CURSOS SENAI", level=1)
    
    p = doc.add_paragraph("Última Atualização: ")
    p.add_run("[DD/MM/AAAA]").bold = True
    
    # Informações Gerais
    doc.add_heading("INFORMAÇÕES GERAIS DA UNIDADE", level=2)
    doc.add_paragraph("Nome da Unidade: SENAI Botucatu", style='List Bullet')
    doc.add_paragraph("Horário de Atendimento: Segunda a Sexta, das 08h às 21h", style='List Bullet')
    doc.add_paragraph("Contato: (14) 3811-0000 / contato@sp.senai.br", style='List Bullet')
    
    # Estrutura do Curso
    doc.add_heading("CURSO 1: [NOME DO CURSO]", level=2)
    doc.add_paragraph("Categoria: [Ex: Técnico / Qualificação]", style='List Bullet')
    doc.add_paragraph("Carga Horária: [Ex: 1200 horas]", style='List Bullet')
    doc.add_paragraph("Modalidade: [Ex: Presencial / EAD]", style='List Bullet')
    doc.add_paragraph("Investimento: [Ex: Gratuito / 10x de R$ 250,00]", style='List Bullet')
    doc.add_paragraph("Pré-requisitos: [Ex: 16 anos completos, Ensino Médio Cursando]", style='List Bullet')
    
    # Conteúdo Programático
    doc.add_heading("Conteúdo Programático:", level=3)
    doc.add_paragraph("Módulo 1: [Nome do Módulo] - [Breve resumo das competências]", style='List Bullet')
    doc.add_paragraph("Módulo 2: [Nome do Módulo] - [Breve resumo das competências]", style='List Bullet')
    
    # Salva o arquivo em um buffer de memória (BytesIO)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0) # Retorna o ponteiro para o início do arquivo
    
    return buffer